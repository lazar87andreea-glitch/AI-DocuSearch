import os
from typing import List


def extract_text_from_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise RuntimeError("pypdf is required for PDF extraction. Install with `pip install pypdf`") from exc

    reader = PdfReader(path)
    texts: List[str] = []
    for page in reader.pages:
        texts.append(page.extract_text() or "")
    return "\n".join(texts).strip()


def extract_text_from_pdf_ocr(path: str, poppler_path: str | None = None) -> str:
    """Fallback OCR extraction for scanned PDFs using pdf2image + pytesseract.

    Optional parameters:
    - `poppler_path`: path to Poppler `bin` directory; if not provided the
      environment variable `POPPLER_PATH` will be used if present.

    This function raises a RuntimeError with installation hints if required
    libraries or system binaries are not available.
    """
    try:
        from pdf2image import convert_from_path
    except Exception as exc:
        raise RuntimeError(
            "pdf2image is required for OCR fallback. Install with `pip install pdf2image pillow pytesseract` and ensure poppler is installed on your system (see https://pdf2image.readthedocs.io/en/latest/)."
        ) from exc

    try:
        import pytesseract
    except Exception as exc:
        raise RuntimeError(
            "pytesseract is required for OCR fallback. Install with `pip install pytesseract` and ensure Tesseract OCR is installed on your system."
        ) from exc

    # Allow overriding the tesseract executable via TESSERACT_CMD env var
    env_tesseract = os.environ.get("TESSERACT_CMD")
    if env_tesseract:
        pytesseract.pytesseract.tesseract_cmd = env_tesseract

    # Determine poppler path (explicit argument overrides environment)
    env_poppler = os.environ.get("POPPLER_PATH")
    poppler_arg = poppler_path or env_poppler

    # Convert PDF pages to images. If poppler path is provided, pass it through.
    if poppler_arg:
        images = convert_from_path(path, dpi=300, poppler_path=poppler_arg)
    else:
        images = convert_from_path(path, dpi=300)

    page_texts: List[str] = []
    for img in images:
        try:
            text = pytesseract.image_to_string(img, lang=None)
        except Exception:
            text = ""
        if text:
            page_texts.append(text)
    return "\n".join(page_texts).strip()


def extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for DOCX extraction. Install with `pip install python-docx`") from exc

    document = Document(path)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_text_from_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path: str) -> str:
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        # Try normal PDF extraction first
        text = extract_text_from_pdf(path)
        if text and text.strip():
            return text

        # If no text found, attempt OCR fallback. Allows using POPPLER_PATH env var.
        try:
            poppler_env = os.environ.get("POPPLER_PATH")
            return extract_text_from_pdf_ocr(path, poppler_path=poppler_env)
        except Exception:
            # If OCR fails, return whatever (possibly empty) text we have
            return text

    if ext == ".docx":
        return extract_text_from_docx(path)
    if ext == ".doc":
        raise RuntimeError("The .doc format is not supported directly. Convert to .docx or plain text before ingestion.")

    return extract_text_from_text_file(path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python src/ingest.py <file>")
        sys.exit(1)

    file_path = sys.argv[1]
    try:
        content = extract_text(file_path)
        if content:
            print(content)
        else:
            print("No text was extracted from the file.")
    except Exception as exc:
        print(f"Error: {exc}")
        sys.exit(1)
