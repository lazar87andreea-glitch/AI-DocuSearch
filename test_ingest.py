import os
import sys
import importlib.util
import time
from pathlib import Path
from types import ModuleType
from unittest.mock import Mock, patch

from src.ingest import (
    extract_text,
    extract_text_from_pdf,
    extract_text_from_pdf_ocr,
)
from src.pipeline import build_pipeline_from_text
from src.preprocess import chunk_text
from src.upload_storage import (
    UPLOAD_PREFIX,
    UPLOAD_TEMP_DIR,
    cleanup_stale_uploads,
    temporary_upload,
)


def test_extract_text_plain_text():
    path = "temp_ingest_test.txt"
    content = "Hello world\nThis is a test."
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

    try:
        result = extract_text(path)
        assert result == content, f"Expected {content!r}, got {result!r}"
    finally:
        os.remove(path)


def test_extract_text_file_not_found():
    missing_path = "this_file_does_not_exist.txt"
    try:
        extract_text(missing_path)
        assert False, "Expected FileNotFoundError"
    except FileNotFoundError:
        pass


def test_extract_text_docx_support():
    spec = importlib.util.find_spec("docx")
    if spec is None:
        print("SKIPPED: python-docx is not installed")
        return

    from docx import Document

    path = "temp_ingest_test.docx"
    doc = Document()
    doc.add_paragraph("Hello DOCX")
    doc.add_paragraph("This is a document test.")
    doc.save(path)

    try:
        result = extract_text(path)
        assert "Hello DOCX" in result, "DOCX paragraph missing"
        assert "This is a document test." in result, "DOCX paragraph missing"
    finally:
        os.remove(path)


def test_extract_text_doc_unsupported():
    path = "temp_ingest_test.doc"
    with open(path, "w", encoding="utf-8") as f:
        f.write("Dummy legacy doc content")

    try:
        try:
            extract_text(path)
            assert False, "Expected RuntimeError for .doc support"
        except RuntimeError as exc:
            assert ".doc format is not supported" in str(exc)
    finally:
        os.remove(path)


def test_searchable_pdf_extraction_preserves_page_numbers():
    pages = []
    for text in ["Cover", "", "Chapter one"]:
        page = Mock()
        page.extract_text.return_value = text
        pages.append(page)

    with patch("pypdf.PdfReader", return_value=Mock(pages=pages)):
        result = extract_text_from_pdf("example.pdf")

    assert result == (
        "[PDF_PAGE:1]\nCover\n\n"
        "[PDF_PAGE:2]\n\n\n"
        "[PDF_PAGE:3]\nChapter one"
    )


def test_ocr_pdf_extraction_preserves_empty_page_numbers():
    images = [object(), object(), object()]
    pdf2image = ModuleType("pdf2image")
    pdf2image.convert_from_path = Mock(return_value=images)
    pytesseract = ModuleType("pytesseract")
    pytesseract.pytesseract = Mock()
    pytesseract.image_to_string = Mock(
        side_effect=["Cover", RuntimeError(), "Body"]
    )

    with patch.dict(
        sys.modules,
        {"pdf2image": pdf2image, "pytesseract": pytesseract},
    ):
        result = extract_text_from_pdf_ocr("example.pdf")

    assert "[PDF_PAGE:1]\nCover" in result
    assert result.index("[PDF_PAGE:1]") < result.index("[PDF_PAGE:2]")
    assert result.index("[PDF_PAGE:2]") < result.index("[PDF_PAGE:3]")
    assert result.endswith("[PDF_PAGE:3]\nBody")


def test_marker_only_pdf_extraction_triggers_ocr():
    path = Path("temp_scanned_test.pdf")
    path.touch()
    try:
        with (
            patch(
                "src.ingest.extract_text_from_pdf",
                return_value="[PDF_PAGE:1]\n\n[PDF_PAGE:2]",
            ),
            patch(
                "src.ingest.extract_text_from_pdf_ocr",
                return_value="[PDF_PAGE:1]\nScanned content",
            ) as ocr,
        ):
            result = extract_text(str(path))

        assert result == "[PDF_PAGE:1]\nScanned content"
        ocr.assert_called_once()
    finally:
        path.unlink(missing_ok=True)


def test_pdf_chunking_never_crosses_page_boundaries():
    text = (
        "[PDF_PAGE:1]\n" + ("First page. " * 80) + "\n\n"
        "[PDF_PAGE:2]\n" + ("Second page. " * 80)
    )

    chunks = chunk_text(text, chunk_size=120, overlap=20)

    assert len(chunks) > 2
    assert all(chunk.startswith("[PDF_PAGE:") for chunk in chunks)
    assert all(
        not ("First page." in chunk and "Second page." in chunk)
        for chunk in chunks
    )
    assert any(chunk.startswith("[PDF_PAGE:1]") for chunk in chunks)
    assert any(chunk.startswith("[PDF_PAGE:2]") for chunk in chunks)


def test_temporary_upload_cleanup_after_success():
    upload_path = ""
    with temporary_upload("example.txt", b"temporary content") as path:
        upload_path = path
        assert os.path.exists(path)
        assert extract_text(path) == "temporary content"

    assert not os.path.exists(upload_path)


def test_temporary_upload_cleanup_after_failure():
    upload_path = ""
    try:
        with temporary_upload("broken.txt", b"temporary content") as path:
            upload_path = path
            raise RuntimeError("simulated extraction failure")
    except RuntimeError as exc:
        assert str(exc) == "simulated extraction failure"

    assert not os.path.exists(upload_path)


def test_pipeline_remains_usable_after_upload_cleanup():
    with temporary_upload("example.txt", b"The contract starts on Monday.") as path:
        upload_path = path
        document_text = extract_text(path)

    pipeline = build_pipeline_from_text(document_text, use_embeddings=False)

    assert not os.path.exists(upload_path)
    assert pipeline["chunks"] == ["The contract starts on Monday."]


def test_stale_upload_cleanup_is_scoped():
    UPLOAD_TEMP_DIR.mkdir(parents=True, exist_ok=True)
    stale_upload = UPLOAD_TEMP_DIR / f"{UPLOAD_PREFIX}stale.txt"
    current_upload = UPLOAD_TEMP_DIR / f"{UPLOAD_PREFIX}current.txt"
    unrelated_file = Path(UPLOAD_TEMP_DIR.parent) / "unrelated_temp_file.txt"

    stale_upload.write_text("stale", encoding="utf-8")
    current_upload.write_text("current", encoding="utf-8")
    unrelated_file.write_text("unrelated", encoding="utf-8")
    old_time = time.time() - 7200
    os.utime(stale_upload, (old_time, old_time))

    try:
        cleanup_stale_uploads(max_age_seconds=3600)

        assert not stale_upload.exists()
        assert current_upload.exists()
        assert unrelated_file.exists()
    finally:
        current_upload.unlink(missing_ok=True)
        unrelated_file.unlink(missing_ok=True)


def run_tests():
    test_extract_text_plain_text()
    test_extract_text_file_not_found()
    test_extract_text_docx_support()
    test_extract_text_doc_unsupported()
    test_searchable_pdf_extraction_preserves_page_numbers()
    test_ocr_pdf_extraction_preserves_empty_page_numbers()
    test_marker_only_pdf_extraction_triggers_ocr()
    test_pdf_chunking_never_crosses_page_boundaries()
    test_temporary_upload_cleanup_after_success()
    test_temporary_upload_cleanup_after_failure()
    test_pipeline_remains_usable_after_upload_cleanup()
    test_stale_upload_cleanup_is_scoped()
    print("All ingestion tests completed.")


def run_extract(path: str):
    content = extract_text(path)
    print(content)


if __name__ == "__main__":
    if len(sys.argv) == 1:
        run_tests()
    elif len(sys.argv) == 2:
        run_extract(sys.argv[1])
    else:
        print("Usage:\n  python test_ingest.py           # run ingestion tests\n  python test_ingest.py <file>    # extract text from a file")
        sys.exit(1)
